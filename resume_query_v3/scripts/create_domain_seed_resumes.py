from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCUMENTS_DIR = REPO_ROOT / "documents"
V3_RESUME_DIR = REPO_ROOT / "resume"


DOMAIN_RESUMES: list[dict[str, Any]] = [
    {
        "file": "林妍-用户运营专员.docx",
        "name": "林妍",
        "role": "用户运营专员",
        "phone": "13800010001",
        "email": "lin.yan@example.com",
        "city": "上海",
        "years": "4年",
        "intent": "用户运营 / 会员运营",
        "summary": "长期负责用户分层、留存召回、会员活动和运营复盘，熟悉 SQL、Excel、用户画像和生命周期运营。",
        "skills": "用户运营、会员运营、用户分层、留存召回、A/B测试、SQL、Excel、数据复盘",
        "projects": [
            ("会员留存提升项目", "SQL、Excel、A/B测试", "基于用户活跃、消费频次和会员等级做分层触达，设计积分激励与召回策略。", "会员月留存率提升12%，沉淀用户运营复盘模板。"),
            ("新用户激活运营项目", "用户画像、短信平台、数据看板", "设计新用户7日激活路径，联动产品优化新人任务和首单权益。", "新用户首周激活率提升18%，首单转化率提升9%。"),
        ],
        "work": ["2022.03 - 2026.04  星橙生活  用户运营专员：负责用户分层、留存召回和会员活动运营。"],
        "education": ["华东师范大学  市场营销本科  2018.09 - 2022.06"],
    },
    {
        "file": "许泽-增长运营经理.docx",
        "name": "许泽",
        "role": "增长运营经理",
        "phone": "13800010002",
        "email": "xu.ze@example.com",
        "city": "杭州",
        "years": "6年",
        "intent": "增长运营 / 用户增长",
        "summary": "负责增长实验、渠道投放、转化漏斗和运营数据分析，能够把增长目标拆成可执行实验。",
        "skills": "增长运营、渠道运营、转化漏斗、A/B测试、SQL、Python、Tableau、运营复盘",
        "projects": [
            ("增长实验中台项目", "SQL、Python、Tableau", "搭建注册、激活、转化和复购漏斗，按渠道和用户群设计增长实验。", "月新增有效用户提升28%，投放获客成本下降16%。"),
            ("内容转化优化项目", "A/B测试、数据看板、埋点分析", "分析内容曝光、点击和转化链路，优化落地页和权益展示。", "核心落地页转化率提升11%，形成周度增长复盘机制。"),
        ],
        "work": ["2020.08 - 2026.04  轻舟科技  增长运营经理：负责用户增长、渠道投放和转化分析。"],
        "education": ["浙江大学  信息管理本科  2016.09 - 2020.06"],
    },
    {
        "file": "唐薇-活动运营经理.docx",
        "name": "唐薇",
        "role": "活动运营经理",
        "phone": "13800010003",
        "email": "tang.wei@example.com",
        "city": "北京",
        "years": "5年",
        "intent": "活动运营 / 品牌活动运营",
        "summary": "擅长活动策划、资源协调、活动数据复盘和用户转化链路设计。",
        "skills": "活动运营、内容运营、活动策划、用户转化、预算管理、Excel、飞书、数据复盘",
        "projects": [
            ("周年大促活动运营项目", "Excel、飞书、BI看板", "负责活动节奏、权益配置、素材排期和跨部门协调，跟踪曝光、点击、转化和复购。", "活动 GMV 提升35%，老客复购率提升14%。"),
            ("线下社群活动项目", "社群运营、问卷、Excel", "组织城市用户沙龙，收集需求反馈并转化为后续内容选题。", "活动满意度92%，沉淀高价值用户线索800+。"),
        ],
        "work": ["2021.01 - 2026.04  青橙消费  活动运营经理：负责大型营销活动和运营复盘。"],
        "education": ["中国传媒大学  广告学本科  2017.09 - 2021.06"],
    },
    {
        "file": "韩子墨-商家运营主管.docx",
        "name": "韩子墨",
        "role": "商家运营主管",
        "phone": "13800010004",
        "email": "han.zimo@example.com",
        "city": "成都",
        "years": "7年",
        "intent": "商家运营 / 平台运营",
        "summary": "负责平台商家分层、经营诊断、活动报名和服务质量提升，熟悉供给侧运营。",
        "skills": "商家运营、平台运营、经营诊断、商家分层、活动运营、SQL、Excel、SLA管理",
        "projects": [
            ("商家分层经营项目", "SQL、Excel、CRM", "按销售额、履约、评分和活跃度拆分商家层级，制定差异化运营动作。", "核心商家销售额提升22%，低活跃商家唤醒率提升19%。"),
            ("商家履约质量提升项目", "SLA看板、数据复盘", "定位发货延迟和售后问题商家，建立预警和辅导机制。", "履约投诉率下降26%，平台平均评分提升0.3。"),
        ],
        "work": ["2019.05 - 2026.04  西岭电商  商家运营主管：负责商家分层、经营诊断和履约质量提升。"],
        "education": ["西南财经大学  电子商务本科  2015.09 - 2019.06"],
    },
    {
        "file": "顾婉清-会员运营负责人.docx",
        "name": "顾婉清",
        "role": "会员运营负责人",
        "phone": "13800010005",
        "email": "gu.wanqing@example.com",
        "city": "广州",
        "years": "8年",
        "intent": "会员运营 / 用户运营负责人",
        "summary": "负责会员体系、权益设计、用户生命周期和精细化触达，关注复购与客单价提升。",
        "skills": "会员运营、用户生命周期、权益体系、CRM、用户分层、SQL、Tableau、数据复盘",
        "projects": [
            ("会员权益体系升级项目", "CRM、SQL、Tableau", "重构会员等级、积分和专属权益，按用户价值设计差异化触达。", "会员复购率提升17%，高价值会员客单价提升13%。"),
            ("沉睡会员召回项目", "短信平台、A/B测试、Excel", "按沉睡周期和历史偏好设计召回包和内容策略。", "沉睡会员召回率提升21%，召回成本下降10%。"),
        ],
        "work": ["2018.02 - 2026.04  岭南零售  会员运营负责人：负责会员体系、CRM和精细化运营。"],
        "education": ["暨南大学  工商管理本科  2014.09 - 2018.06"],
    },
    {
        "file": "秦昊-新能源项目经理.docx",
        "name": "秦昊",
        "role": "新能源项目经理",
        "phone": "13800030001",
        "email": "qin.hao@example.com",
        "city": "上海",
        "years": "8年",
        "intent": "新能源项目经理 / 光伏项目管理",
        "summary": "负责光伏项目立项、招采、EPC交付、并网验收和项目成本复盘，熟悉新能源项目管理。",
        "skills": "新能源项目、光伏电站、EPC管理、招投标、项目进度、成本控制、Excel、Power BI",
        "projects": [
            ("分布式光伏电站建设项目", "Excel、Power BI、项目管理系统", "主导园区分布式光伏项目立项、招采、施工进度和并网验收。", "项目容量42MW，按期完成并网，施工返工率降低18%。"),
            ("新能源项目成本复盘体系", "Excel、SQL、PowerPoint", "搭建设备采购、施工、并网和运维成本复盘模型。", "支持8个新能源项目预算复盘，设备采购成本平均下降6%。"),
        ],
        "work": ["2021.03 - 2026.04  华曜新能源  新能源项目经理：负责光伏项目交付、成本控制和跨部门协调。"],
        "education": ["上海电力大学  电气工程硕士  2015.09 - 2018.06"],
    },
    {
        "file": "罗雨晴-储能产品运营.docx",
        "name": "罗雨晴",
        "role": "储能产品运营",
        "phone": "13800030002",
        "email": "luo.yuqing@example.com",
        "city": "深圳",
        "years": "5年",
        "intent": "储能产品运营 / 能源解决方案运营",
        "summary": "负责储能产品客户需求、售前方案、交付问题闭环和运行数据看板。",
        "skills": "储能系统、产品运营、客户需求分析、售前方案、数据看板、SQL、Tableau、用户培训",
        "projects": [
            ("工商业储能产品运营体系", "SQL、Tableau、CRM", "负责储能产品线客户需求归因、案例库建设和交付问题闭环。", "售前方案复用率提升45%，支持销售团队签约12个储能项目。"),
            ("储能设备运行数据看板", "SQL、Tableau、Python", "梳理储能设备运行、充放电效率、告警和收益指标。", "看板覆盖30+站点，异常定位时间从2小时降低到20分钟。"),
        ],
        "work": ["2022.01 - 2026.04  星河储能  储能产品运营：负责储能解决方案、客户运营和数据看板。"],
        "education": ["华南理工大学  工业工程本科  2015.09 - 2019.06"],
    },
    {
        "file": "何嘉诚-风电运维工程师.docx",
        "name": "何嘉诚",
        "role": "风电运维工程师",
        "phone": "13800030003",
        "email": "he.jiacheng@example.com",
        "city": "乌鲁木齐",
        "years": "6年",
        "intent": "风电运维工程师 / 新能源运维管理",
        "summary": "负责风电场运维、SCADA故障诊断、备件管理和设备可靠性分析。",
        "skills": "风电运维、SCADA、故障诊断、备件管理、设备巡检、Python、Excel、可靠性分析",
        "projects": [
            ("风电场故障诊断与停机分析项目", "SCADA、Python、Excel", "基于 SCADA 数据整理风机告警、停机和功率曲线。", "非计划停机时长下降21%，月度运维报告自动化生成。"),
            ("风机备件库存优化项目", "Excel、SQL、ERP", "建立备件消耗、故障频率和采购周期模型。", "关键备件缺货次数降低40%，库存资金占用下降12%。"),
        ],
        "work": ["2020.04 - 2026.04  北疆风电  风电运维工程师：负责风电场运维、故障诊断和备件管理。"],
        "education": ["华北电力大学  新能源科学与工程本科  2014.09 - 2018.06"],
    },
    {
        "file": "蒋安然-电力交易分析师.docx",
        "name": "蒋安然",
        "role": "电力交易分析师",
        "phone": "13800030004",
        "email": "jiang.anran@example.com",
        "city": "广州",
        "years": "4年",
        "intent": "电力交易分析师 / 能源数据分析",
        "summary": "负责电力现货交易、负荷预测、报价策略、绿电交易和能源数据分析。",
        "skills": "电力交易、现货市场、负荷预测、SQL、Python、数据建模、报价策略、绿电交易",
        "projects": [
            ("电力现货报价策略分析项目", "Python、SQL、Pandas", "清洗负荷、气象、历史出清价格和机组约束数据，构建报价辅助模型。", "报价模拟准确率提升到82%，交易数据整理时间减少60%。"),
            ("绿电交易客户收益测算工具", "Excel、Python、Power BI", "搭建绿电采购成本和碳减排收益测算模板。", "支持20+客户完成绿电方案测算，方案制作时间缩短75%。"),
        ],
        "work": ["2022.06 - 2026.04  粤能售电  电力交易分析师：负责现货交易数据分析和报价策略支持。"],
        "education": ["中山大学  应用统计硕士  2018.09 - 2020.06"],
    },
    {
        "file": "叶晨-碳资产项目顾问.docx",
        "name": "叶晨",
        "role": "碳资产项目顾问",
        "phone": "13800030005",
        "email": "ye.chen@example.com",
        "city": "北京",
        "years": "5年",
        "intent": "碳资产项目顾问 / 能源低碳咨询",
        "summary": "负责碳排放核算、能源审计、碳资产管理、ESG数据治理和低碳项目交付。",
        "skills": "碳资产管理、碳排放核算、能源审计、ESG报告、Excel、Python、政策研究、项目交付",
        "projects": [
            ("制造企业碳排放核算项目", "Excel、Python、PowerPoint", "负责能源消耗数据收集、排放因子匹配和碳排放核算模型搭建。", "完成5家工厂碳盘查，发现节能改造机会12项。"),
            ("新能源企业 ESG 数据治理项目", "Excel、飞书、数据看板", "梳理新能源企业环境、能源、供应链和治理指标。", "ESG 数据填报错误率下降35%，报告编制周期缩短40%。"),
        ],
        "work": ["2021.08 - 2026.04  青岚低碳咨询  碳资产项目顾问：负责碳盘查、能源审计和低碳项目交付。"],
        "education": ["北京化工大学  环境工程硕士  2016.09 - 2019.06"],
    },
    {
        "file": "陈思远-金融产品经理.docx",
        "name": "陈思远",
        "role": "金融产品经理",
        "phone": "13800020001",
        "email": "chen.siyuan@example.com",
        "city": "上海",
        "years": "7年",
        "intent": "金融产品经理 / 财富管理产品",
        "summary": "负责金融产品需求、基金交易、支付链路、合规风控和数据指标体系。",
        "skills": "金融产品、基金交易、支付、风控、需求分析、SQL、Axure、数据看板",
        "projects": [
            ("基金交易产品改版项目", "Axure、SQL、埋点分析", "负责申购、赎回、持仓和收益展示链路优化，联动合规校验。", "交易转化率提升15%，客服咨询量下降22%。"),
            ("金融产品风险提示项目", "风控规则、SQL、产品运营", "设计风险测评、适当性提示和高风险产品确认流程。", "高风险误购投诉下降30%，合规审核一次通过。"),
        ],
        "work": ["2019.06 - 2026.04  恒川财富  金融产品经理：负责基金交易和财富管理产品。"],
        "education": ["上海财经大学  金融学本科  2015.09 - 2019.06"],
    },
    {
        "file": "郭明昊-金融数据分析师.docx",
        "name": "郭明昊",
        "role": "金融数据分析师",
        "phone": "13800020002",
        "email": "guo.minghao@example.com",
        "city": "北京",
        "years": "5年",
        "intent": "金融数据分析 / 风险分析",
        "summary": "负责信贷数据分析、风控报表、资产质量监控和金融业务指标体系。",
        "skills": "金融数据分析、信贷风控、SQL、Python、Pandas、风险指标、Power BI、资产质量",
        "projects": [
            ("信贷资产质量监控项目", "SQL、Python、Power BI", "搭建逾期率、迁徙率和坏账预测监控看板。", "风险预警提前7天，人工报表时间减少70%。"),
            ("授信策略效果复盘项目", "Python、Pandas、统计分析", "按客群、渠道和额度段分析授信策略表现。", "识别高风险渠道3个，推动策略调整后逾期率下降1.8个百分点。"),
        ],
        "work": ["2021.02 - 2026.04  京融科技  金融数据分析师：负责信贷风控和资产质量分析。"],
        "education": ["中央财经大学  统计学硕士  2018.09 - 2021.06"],
    },
    {
        "file": "孙可欣-风控运营经理.docx",
        "name": "孙可欣",
        "role": "风控运营经理",
        "phone": "13800020003",
        "email": "sun.kexin@example.com",
        "city": "深圳",
        "years": "6年",
        "intent": "风控运营 / 金融风控",
        "summary": "负责反欺诈策略、信贷风控运营、风险名单、规则上线和策略复盘。",
        "skills": "风控运营、反欺诈、信贷风控、策略规则、SQL、Python、名单管理、风险复盘",
        "projects": [
            ("反欺诈策略运营项目", "SQL、Python、规则引擎", "分析设备、行为、交易和黑名单特征，维护反欺诈规则。", "欺诈损失率下降24%，误杀率下降8%。"),
            ("贷后风险预警项目", "SQL、风险模型、数据看板", "搭建逾期预警、催收分层和风险名单策略。", "M1逾期率下降1.5个百分点，催收触达效率提升20%。"),
        ],
        "work": ["2020.04 - 2026.04  南湾金融  风控运营经理：负责反欺诈和贷后风险运营。"],
        "education": ["深圳大学  金融工程本科  2016.09 - 2020.06"],
    },
    {
        "file": "刘嘉宁-量化研究助理.docx",
        "name": "刘嘉宁",
        "role": "量化研究助理",
        "phone": "13800020004",
        "email": "liu.jianing@example.com",
        "city": "上海",
        "years": "3年",
        "intent": "量化研究 / 金融策略分析",
        "summary": "负责股票因子研究、回测、组合分析和金融市场数据处理。",
        "skills": "量化研究、股票交易、因子模型、回测、Python、Pandas、NumPy、金融数据",
        "projects": [
            ("多因子选股回测项目", "Python、Pandas、NumPy", "清洗股票行情、财务和行业数据，构建估值、动量和质量因子。", "回测年化收益12.6%，最大回撤控制在8.4%。"),
            ("交易信号监控工具", "Python、SQL、定时任务", "搭建因子信号、调仓列表和异常数据检查流程。", "每日策略生成时间从2小时缩短到20分钟。"),
        ],
        "work": ["2023.01 - 2026.04  沪上量化  量化研究助理：负责因子研究和回测分析。"],
        "education": ["复旦大学  金融工程硕士  2020.09 - 2023.06"],
    },
    {
        "file": "周雨桐-投研分析师.docx",
        "name": "周雨桐",
        "role": "投研分析师",
        "phone": "13800020005",
        "email": "zhou.yutong@example.com",
        "city": "北京",
        "years": "4年",
        "intent": "投研分析 / 行业研究",
        "summary": "负责基金投研、行业研究、财务建模、估值分析和投资报告撰写。",
        "skills": "投研分析、基金研究、财务建模、估值分析、Wind、Excel、Python、行业研究",
        "projects": [
            ("新能源行业投研项目", "Wind、Excel、Python", "跟踪新能源产业链公司财务、估值和政策变化，撰写行业研究报告。", "覆盖12家公司，报告被投委会用于仓位调整讨论。"),
            ("基金组合归因分析项目", "Excel、Python、Barra风格因子", "分析基金组合收益来源、行业暴露和风格漂移。", "帮助组合经理识别超额收益来源和回撤风险。"),
        ],
        "work": ["2022.05 - 2026.04  北辰资管  投研分析师：负责行业研究和基金组合分析。"],
        "education": ["对外经济贸易大学  金融硕士  2019.09 - 2022.06"],
    },
]


def main() -> int:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    V3_RESUME_DIR.mkdir(parents=True, exist_ok=True)
    for resume in DOMAIN_RESUMES:
        payload = _build_resume_bytes(resume)
        _write_resume_bytes(V3_RESUME_DIR / str(resume["file"]), payload)
        _write_resume_bytes(DOCUMENTS_DIR / str(resume["file"]), payload)
    print(f"created_or_updated={len(DOMAIN_RESUMES)}")
    for resume in DOMAIN_RESUMES:
        print(str(DOCUMENTS_DIR / str(resume["file"])))
        print(str(V3_RESUME_DIR / str(resume["file"])))
    return 0


def _build_resume_bytes(resume: dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(str(resume["name"]), level=1)
    doc.add_paragraph(str(resume["role"]))
    doc.add_heading("基本信息", level=2)
    doc.add_paragraph(f"手机：{resume['phone']}")
    doc.add_paragraph(f"邮箱：{resume['email']}")
    doc.add_paragraph(f"现居：{resume['city']}")
    doc.add_paragraph(f"工作年限：{resume['years']}")
    doc.add_paragraph(f"求职意向：{resume['intent']}")
    doc.add_heading("个人总结", level=2)
    doc.add_paragraph(str(resume["summary"]))
    doc.add_heading("专业技能", level=2)
    doc.add_paragraph(str(resume["skills"]))
    doc.add_heading("项目经历", level=2)
    for index, (name, stack, duty, result) in enumerate(resume["projects"], start=1):
        doc.add_paragraph(f"项目{index}：{name}")
        doc.add_paragraph(f"技术栈：{stack}")
        doc.add_paragraph(f"职责：{duty}")
        doc.add_paragraph(f"成果：{result}")
    doc.add_heading("工作经历", level=2)
    for item in resume["work"]:
        doc.add_paragraph(str(item))
    doc.add_heading("教育背景", level=2)
    for item in resume["education"]:
        doc.add_paragraph(str(item))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _write_resume_bytes(path: Path, payload: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(payload)


if __name__ == "__main__":
    raise SystemExit(main())
