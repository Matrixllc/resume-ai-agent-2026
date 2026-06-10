from __future__ import annotations

import html
from pathlib import Path
from urllib.parse import quote

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, HTMLResponse

from resume_query_tools.config import get_tools_config
from resume_query_tools.stores.sql_reader import ResumeSqlReader
from resume_query_tools import (
    build_candidate_summary_context,
    generate_candidate_summary,
    get_candidate_profile_display,
    get_classified_projects_display,
    list_candidates_display,
)

router = APIRouter(prefix="/candidates", tags=["candidates"])

ALLOWED_RESUME_EXTENSIONS = {".pdf", ".doc", ".docx"}
RESUME_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@router.get("")
def candidate_list() -> dict:
    return list_candidates_display().model_dump()


@router.get("/{resume_identity}")
def candidate_detail(resume_identity: str) -> dict:
    try:
        return get_candidate_profile_display(resume_identity).model_dump()
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error


@router.get("/{resume_identity}/resume-document")
def candidate_resume_document(resume_identity: str) -> dict:
    document = _resolve_resume_document(resume_identity, require_available=False)
    document.pop("source_path", None)
    if not document["source_available"]:
        return document
    return document


@router.get("/{resume_identity}/resume-document/preview")
def candidate_resume_document_preview(resume_identity: str):
    document = _resolve_resume_document(resume_identity, require_available=True)
    source_path = Path(document["source_path"])
    if document["preview_kind"] == "pdf":
        return FileResponse(
            source_path,
            media_type=document["mime_type"],
            filename=document["file_name"],
            headers={"Content-Disposition": f"inline; filename*=UTF-8''{quote(document['file_name'])}"},
        )
    if document["preview_kind"] == "html":
        preview_html = _resume_to_safe_html(source_path)
        return HTMLResponse(preview_html)
    raise HTTPException(status_code=415, detail="resume document cannot be previewed")


@router.get("/{resume_identity}/resume-document/file")
def candidate_resume_document_file(resume_identity: str):
    document = _resolve_resume_document(resume_identity, require_available=True)
    source_path = Path(document["source_path"])
    return FileResponse(
        source_path,
        media_type=document["mime_type"],
        filename=document["file_name"],
    )


@router.get("/{resume_identity}/projects")
def candidate_projects(resume_identity: str) -> dict:
    try:
        return get_classified_projects_display(resume_identity).model_dump()
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error


@router.get("/{resume_identity}/work-experiences")
def candidate_work_experiences(resume_identity: str) -> dict:
    try:
        detail = get_candidate_profile_display(resume_identity)
        return {"resume_identity": resume_identity, "work_experiences": [item.model_dump() for item in detail.work_experiences]}
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error


def _resolve_resume_document(resume_identity: str, *, require_available: bool) -> dict:
    config = get_tools_config()
    reader = ResumeSqlReader(config["paths"]["structured_store_file"])
    candidate = reader.get_candidate(resume_identity)
    if not candidate:
        raise HTTPException(status_code=404, detail=f"candidate not found: {resume_identity}")

    source_raw = str(candidate.get("source_path", "") or "").strip()
    file_name = Path(source_raw).name if source_raw else str(candidate.get("file_name", "") or "")
    source_path = Path(source_raw).expanduser().resolve() if source_raw else None
    extension = source_path.suffix.lower() if source_path else Path(file_name).suffix.lower()
    is_allowed = extension in ALLOWED_RESUME_EXTENSIONS
    source_available = bool(source_path and source_path.exists() and source_path.is_file() and is_allowed)

    if require_available and not source_available:
        if not source_path or not source_path.exists():
            raise HTTPException(status_code=404, detail="resume source file not found")
        if not is_allowed:
            raise HTTPException(status_code=415, detail="resume source file type is not supported")
        raise HTTPException(status_code=404, detail="resume source file is not available")

    encoded_identity = quote(resume_identity, safe="")
    preview_kind = ""
    if source_available:
        preview_kind = "pdf" if extension == ".pdf" else "html"
    return {
        "resume_identity": resume_identity,
        "file_name": file_name,
        "extension": extension,
        "mime_type": RESUME_MIME_TYPES.get(extension, "application/octet-stream"),
        "preview_kind": preview_kind,
        "preview_url": f"/candidates/{encoded_identity}/resume-document/preview" if source_available else "",
        "download_url": f"/candidates/{encoded_identity}/resume-document/file" if source_available else "",
        "source_available": source_available,
        "source_path": str(source_path) if source_path else "",
    }


def _resume_to_safe_html(source_path: Path) -> str:
    try:
        paragraphs = _extract_resume_paragraphs(source_path)
        if not paragraphs:
            raise ValueError("empty document")
        body = "\n".join(f"<p>{html.escape(line)}</p>" for line in paragraphs)
    except Exception as error:
        body = (
            "<div class=\"notice\">"
            f"无法预览该文件，可下载原文件查看。原因：{html.escape(type(error).__name__)}"
            "</div>"
        )
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    body {{
      margin: 0;
      padding: 24px;
      color: #1e293b;
      background: #ffffff;
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      font-size: 14px;
      line-height: 1.85;
    }}
    p {{
      margin: 0 0 10px;
      white-space: pre-wrap;
      overflow-wrap: anywhere;
    }}
    .notice {{
      border: 1px solid #fde68a;
      background: #fffbeb;
      color: #92400e;
      border-radius: 8px;
      padding: 12px;
    }}
  </style>
</head>
<body>{body}</body>
</html>"""


def _extract_resume_paragraphs(source_path: Path) -> list[str]:
    suffix = source_path.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        document = Document(str(source_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return paragraphs

    if suffix == ".doc":
        try:
            from docling.document_converter import DocumentConverter  # type: ignore

            result = DocumentConverter().convert(str(source_path))
            text = str(result.document.export_to_markdown() or "")
            return [line.strip() for line in text.splitlines() if line.strip()]
        except Exception:
            text = source_path.read_text(encoding="utf-8", errors="ignore")
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if not lines:
                raise
            return lines
    raise ValueError("unsupported document type")


@router.get("/{resume_identity}/summary-context")
def candidate_summary_context(resume_identity: str) -> dict:
    try:
        return build_candidate_summary_context(resume_identity).model_dump()
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error


@router.post("/{resume_identity}/summary")
def candidate_summary(resume_identity: str) -> dict:
    try:
        summary = generate_candidate_summary(resume_identity)
        return {
            "resume_identity": summary.resume_identity,
            "summary": summary.summary,
            "summary_sections": summary.summary_sections.model_dump(),
            "llm_error": "fallback" if summary.llm_error else "",
        }
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error
